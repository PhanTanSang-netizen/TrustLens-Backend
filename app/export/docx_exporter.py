from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from docx.shared import Inches


def _safe_text(value: Any) -> str:
    if value is None:
        return ""

    return str(value)


def _add_key_value_table(
    document: Document,
    data: dict[str, Any],
) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for key, value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = _safe_text(key)
        row_cells[1].text = _safe_text(value)

    document.add_paragraph("")


def _add_section_heading(
    document: Document,
    title: str,
    level: int = 1,
) -> None:
    document.add_heading(title, level=level)


def _add_submission_info(
    document: Document,
    report_data: dict[str, Any],
) -> None:
    _add_section_heading(document, "1. Thông tin bài nộp", level=1)

    submission = report_data.get("submission") or {}
    file_info = report_data.get("file") or {}

    _add_key_value_table(
        document=document,
        data={
            "Submission ID": submission.get("id"),
            "Assignment ID": submission.get("assignment_id"),
            "Owner Label": submission.get("owner_label"),
            "Submission Status": submission.get("status"),
            "Overall Score": submission.get("overall_score"),
            "Created At": submission.get("created_at"),
            "File Name": file_info.get("original_name"),
            "File Type": file_info.get("mime_type"),
            "File Size Bytes": file_info.get("size_bytes"),
        },
    )


def _add_processing_summary(
    document: Document,
    report_data: dict[str, Any],
) -> None:
    _add_section_heading(document, "2. Tóm tắt xử lý", level=1)

    summary = report_data.get("summary") or {}
    processing = summary.get("processing") or {}
    verification = summary.get("verification") or {}
    score = summary.get("score") or {}

    _add_section_heading(document, "2.1. Trạng thái pipeline", level=2)
    _add_key_value_table(
        document=document,
        data={
            "Submission Status": processing.get("submission_status"),
            "Has Extracted Text": processing.get("has_extracted_text"),
            "Has Reference Section": processing.get("has_reference_section"),
            "Citation Count": processing.get("citation_count"),
            "Metadata Record Count": processing.get("metadata_record_count"),
        },
    )

    _add_section_heading(document, "2.2. Kết quả kiểm chứng metadata", level=2)
    _add_key_value_table(
        document=document,
        data={
            "Total": verification.get("total"),
            "Verified": verification.get("verified"),
            "Basic Metadata Present": verification.get("basic_metadata_present"),
            "Broken": verification.get("broken"),
            "Forbidden": verification.get("forbidden"),
            "Unreachable": verification.get("unreachable"),
            "Not Provided": verification.get("not_provided"),
        },
    )

    _add_section_heading(document, "2.3. Trust Score", level=2)
    _add_key_value_table(
        document=document,
        data={
            "Overall Score": score.get("overall_score"),
            "Trust Level": score.get("trust_level"),
            "Note": score.get("note"),
        },
    )


def _add_reference_section_info(
    document: Document,
    report_data: dict[str, Any],
) -> None:
    _add_section_heading(document, "3. Phần tài liệu tham khảo được nhận diện", level=1)

    reference_section = report_data.get("reference_section")

    if not reference_section:
        document.add_paragraph("Chưa nhận diện được phần tài liệu tham khảo.")
        return

    _add_key_value_table(
        document=document,
        data={
            "Heading": reference_section.get("heading"),
            "Start Index": reference_section.get("start_index"),
            "End Index": reference_section.get("end_index"),
            "Detection Method": reference_section.get("detection_method"),
        },
    )

    document.add_paragraph("Preview:")
    document.add_paragraph(_safe_text(reference_section.get("raw_text_preview")))


def _add_citations(
    document: Document,
    report_data: dict[str, Any],
) -> None:
    _add_section_heading(document, "4. Danh sách citation", level=1)

    citations = report_data.get("citations") or []

    if not citations:
        document.add_paragraph("Chưa có citation được tách.")
        return

    for item in citations:
        citation = item.get("citation") or {}
        metadata = item.get("metadata") or {}
        warnings = item.get("warnings") or []

        sequence_no = citation.get("sequence_no") or ""
        title = citation.get("title") or "Không có tiêu đề"

        _add_section_heading(
            document,
            f"4.{sequence_no}. {title}",
            level=2,
        )

        document.add_paragraph("Citation gốc:")
        document.add_paragraph(_safe_text(citation.get("raw_text")))

        _add_key_value_table(
            document=document,
            data={
                "Detected Style": citation.get("detected_style"),
                "Authors": citation.get("authors"),
                "Title": citation.get("title"),
                "Year": citation.get("year"),
                "DOI": citation.get("doi"),
                "URL": citation.get("url"),
                "Verification Status": metadata.get("verification_status"),
                "Provider": metadata.get("provider"),
                "Confidence Score": metadata.get("confidence_score"),
                "Source URL": metadata.get("source_url"),
                "Trust Level": item.get("trust_level"),
                "Score": item.get("score"),
            },
        )

        if warnings:
            document.add_paragraph("Cảnh báo:")
            for warning in warnings:
                document.add_paragraph(_safe_text(warning), style="List Bullet")
        else:
            document.add_paragraph("Cảnh báo: Không có.")


def build_docx_report_document(
    report_data: dict[str, Any],
) -> Document:
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    document.add_heading("TrustLens - Báo cáo thẩm định tài liệu tham khảo", level=0)

    message = report_data.get("message")
    if message:
        document.add_paragraph(_safe_text(message))

    _add_submission_info(
        document=document,
        report_data=report_data,
    )

    _add_processing_summary(
        document=document,
        report_data=report_data,
    )

    _add_reference_section_info(
        document=document,
        report_data=report_data,
    )

    _add_citations(
        document=document,
        report_data=report_data,
    )

    return document


def export_report_to_docx_file(
    report_data: dict[str, Any],
    output_path: str | Path,
) -> str:
    document = build_docx_report_document(
        report_data=report_data,
    )

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document.save(str(output_path))

    return str(output_path)


def export_report_to_docx_bytes(
    report_data: dict[str, Any],
) -> bytes:
    from io import BytesIO

    document = build_docx_report_document(
        report_data=report_data,
    )

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer.read()


def generate_docx_report_filename(
    submission_id: str,
) -> str:
    return f"trustlens_report_{submission_id}_{uuid4().hex[:8]}.docx"


# Compatibility aliases
def build_docx_report(
    report_data: dict[str, Any],
) -> Document:
    return build_docx_report_document(report_data)


def create_docx_report(
    report_data: dict[str, Any],
) -> bytes:
    return export_report_to_docx_bytes(report_data)