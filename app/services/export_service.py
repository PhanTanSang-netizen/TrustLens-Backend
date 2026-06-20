from datetime import datetime, timezone
from pathlib import Path
from textwrap import wrap
from uuid import UUID

from fastapi import HTTPException, status
from sqlalchemy import select
from sqlalchemy.orm import Session
from app.services.access_control_service import get_accessible_submission_or_404
from app.core.config import settings
from app.models.report import Report, ReportExport


PDF_MIME_TYPE = "application/pdf"
DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)
XLSX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "spreadsheetml.sheet"
)


def _has_model_field(model_class, field_name: str) -> bool:
    return hasattr(model_class, field_name)


def _get_export_root_dir() -> Path:
    upload_dir = getattr(settings, "UPLOAD_DIR", "uploads")
    export_dir = Path(upload_dir) / "exports" / "reports"
    export_dir.mkdir(parents=True, exist_ok=True)
    return export_dir


def _safe_file_part(value: str | None, fallback: str) -> str:
    raw_value = str(value or fallback).strip()
    safe_value = "".join(
        char if char.isalnum() else "_"
        for char in raw_value
    )
    safe_value = "_".join(part for part in safe_value.split("_") if part)
    return (safe_value or fallback)[:80]


def _get_latest_report_by_submission_id(
    db: Session,
    submission_id: UUID,
) -> Report:
    report = db.execute(
        select(Report)
        .where(Report.submission_id == submission_id)
        .order_by(Report.created_at.desc())
    ).scalars().first()

    if report is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={
                "error_code": "REPORT_NOT_FOUND",
                "message": "Không tìm thấy report để export.",
                "details": {
                    "submission_id": str(submission_id),
                },
            },
        )

    return report


def _serialize_report_safely(report: Report) -> dict:
    """
    Lấy dữ liệu report ở dạng dict.

    Không import serialize_report ở đầu file để tránh vòng lặp import.
    """

    try:
        from app.services.report_service import serialize_report

        data = serialize_report(report)

        if isinstance(data, dict):
            return data

        return {
            "raw": str(data),
        }

    except Exception:
        return {
            "report_id": str(report.id),
            "submission_id": str(report.submission_id),
            "summary_json": getattr(report, "summary_json", None),
            "config_version": getattr(report, "config_version", None),
            "disclaimer": getattr(report, "disclaimer", None),
            "note": getattr(report, "note", None),
        }


def _build_report_lines(
    report: Report,
    include_raw_citation: bool = True,
) -> list[str]:
    data = _serialize_report_safely(report)

    lines: list[str] = [
        "TrustLens Academic Reference Report",
        f"Report ID: {report.id}",
        f"Submission ID: {report.submission_id}",
    ]

    overall_score = (
        data.get("overall_score")
        or data.get("report_trust_score")
        or getattr(report, "report_trust_score", None)
    )

    confidence_score = (
        data.get("confidence_score")
        or getattr(report, "confidence_score", None)
    )

    overall_label = (
        data.get("overall_label")
        or getattr(report, "overall_label", None)
    )

    config_version = (
        data.get("config_version")
        or data.get("scoring_config_version")
        or getattr(report, "config_version", None)
        or getattr(report, "scoring_config_version", None)
    )

    if overall_score is not None:
        lines.append(f"Trust Score: {overall_score}")

    if confidence_score is not None:
        try:
            lines.append(f"Confidence: {round(float(confidence_score) * 100)}%")
        except Exception:
            lines.append(f"Confidence: {confidence_score}")

    if overall_label is not None:
        lines.append(f"Label: {overall_label}")

    if config_version is not None:
        lines.append(f"Scoring config: {config_version}")

    lines.append("")
    lines.append("Summary:")

    summary = (
        data.get("summary")
        or data.get("summary_json")
        or getattr(report, "summary_json", None)
    )

    if summary:
        for part in wrap(str(summary), width=100):
            lines.append(part)
    else:
        lines.append("No summary data available.")

    component_summary = (
        data.get("component_summary")
        or getattr(report, "component_summary", None)
    )

    if isinstance(component_summary, dict) and component_summary:
        lines.append("")
        lines.append("Component summary:")

        for key, value in component_summary.items():
            lines.append(f"- {key}: {value}")

    citations = data.get("citations")

    if isinstance(citations, list) and citations:
        lines.append("")
        lines.append("Citation details:")

        for index, citation in enumerate(citations, start=1):
            if not isinstance(citation, dict):
                continue

            normalized_fields = citation.get("normalized_fields")
            scores = citation.get("scores")

            title = citation.get("title")

            if title is None and isinstance(normalized_fields, dict):
                title = normalized_fields.get("title")

            raw_text = citation.get("raw_text")

            score = citation.get("reference_trust_score")

            if score is None and isinstance(scores, dict):
                score = scores.get("reference_trust_score")

            display_text = raw_text if include_raw_citation else title
            display_text = display_text or title or raw_text or "Untitled citation"

            prefix = f"{index}."

            if score is not None:
                prefix = f"{index}. [{score}/100]"

            for part in wrap(f"{prefix} {display_text}", width=100):
                lines.append(part)

    disclaimer = (
        data.get("disclaimer")
        or getattr(report, "disclaimer", None)
        or "TrustLens supports academic review and does not replace lecturer judgement."
    )

    lines.append("")
    lines.append(f"Disclaimer: {disclaimer}")

    return lines


def _escape_pdf_text(text: str) -> str:
    return (
        str(text)
        .replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
        .replace("\r", "")
    )


def _build_minimal_pdf(lines: list[str]) -> bytes:
    stream_lines = [
        "BT",
        "/F1 10 Tf",
        "50 790 Td",
        "14 TL",
    ]

    safe_lines: list[str] = []

    for line in lines:
        if line is None:
            continue

        for wrapped_line in wrap(str(line), width=100):
            safe_lines.append(wrapped_line)

    for line in safe_lines[:52]:
        safe_line = _escape_pdf_text(line[:110])
        stream_lines.append(f"({safe_line}) Tj")
        stream_lines.append("T*")

    stream_lines.append("ET")

    stream = "\n".join(stream_lines).encode("latin-1", errors="replace")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 842] "
            b"/Resources << /Font << /F1 4 0 R >> >> "
            b"/Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        (
            b"<< /Length "
            + str(len(stream)).encode("ascii")
            + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        ),
    ]

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]

    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")

    xref_offset = len(pdf)

    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")

    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))

    pdf.extend(
        f"trailer\n<< /Root 1 0 R /Size {len(objects) + 1} >>\n"
        f"startxref\n{xref_offset}\n"
        "%%EOF\n".encode("ascii")
    )

    return bytes(pdf)


def _write_pdf_file(
    output_path: Path,
    report: Report,
    include_raw_citation: bool = True,
) -> None:
    lines = _build_report_lines(
        report=report,
        include_raw_citation=include_raw_citation,
    )

    output_path.write_bytes(_build_minimal_pdf(lines))


def _write_docx_file(
    output_path: Path,
    report: Report,
    include_raw_citation: bool = True,
) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={
                "error_code": "DOCX_DEPENDENCY_MISSING",
                "message": "Thiếu thư viện python-docx để export DOCX.",
                "details": {
                    "error": str(exc),
                },
            },
        )

    lines = _build_report_lines(
        report=report,
        include_raw_citation=include_raw_citation,
    )

    document = Document()
    document.add_heading("TrustLens Academic Reference Report", level=1)

    for line in lines[1:]:
        if line.strip() == "":
            document.add_paragraph("")
        elif line.endswith(":"):
            document.add_heading(line.replace(":", ""), level=2)
        else:
            document.add_paragraph(line)

    document.save(str(output_path))


def _write_xlsx_file(
    output_path: Path,
    report: Report,
    include_raw_citation: bool = True,
) -> None:
    try:
        from openpyxl import Workbook
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={
                "error_code": "XLSX_DEPENDENCY_MISSING",
                "message": "Thiếu thư viện openpyxl để export XLSX.",
                "details": {
                    "install_command": "pip install openpyxl",
                    "error": str(exc),
                },
            },
        )

    data = _serialize_report_safely(report)

    workbook = Workbook()

    summary_sheet = workbook.active
    summary_sheet.title = "Summary"

    summary_sheet.append(["TrustLens Academic Reference Report"])
    summary_sheet.append(["Report ID", str(report.id)])
    summary_sheet.append(["Submission ID", str(report.submission_id)])

    summary_sheet.append([
        "Overall Score",
        data.get("overall_score") or data.get("report_trust_score"),
    ])

    summary_sheet.append([
        "Overall Label",
        data.get("overall_label"),
    ])

    summary_sheet.append([
        "Config Version",
        data.get("config_version") or data.get("scoring_config_version"),
    ])

    summary_sheet.append([])
    summary_sheet.append(["Summary"])
    summary_sheet.append([
        str(data.get("summary") or data.get("summary_json") or "")
    ])

    citations = data.get("citations")

    if isinstance(citations, list):
        citation_sheet = workbook.create_sheet("Citations")
        citation_sheet.append([
            "No",
            "Raw Text",
            "Title",
            "Score",
            "Status",
            "Warning",
        ])

        for index, citation in enumerate(citations, start=1):
            if not isinstance(citation, dict):
                continue

            normalized_fields = citation.get("normalized_fields")
            scores = citation.get("scores")

            title = citation.get("title")

            if title is None and isinstance(normalized_fields, dict):
                title = normalized_fields.get("title")

            score = citation.get("reference_trust_score")

            if score is None and isinstance(scores, dict):
                score = scores.get("reference_trust_score")

            citation_sheet.append([
                index,
                citation.get("raw_text") if include_raw_citation else "",
                title,
                score,
                citation.get("status"),
                citation.get("warning") or citation.get("message"),
            ])

    workbook.save(str(output_path))


def _record_export_safely(
    db: Session,
    report: Report,
    export_format: str,
    file_name: str,
    stored_path: str,
    mime_type: str,
    status_value: str,
    created_by: UUID | None,
    error_message: str | None = None,
) -> ReportExport:
    path_obj = Path(stored_path)
    size_bytes = path_obj.stat().st_size if path_obj.exists() else None

    data = {
        "report_id": report.id,
    }

    if _has_model_field(ReportExport, "submission_id"):
        data["submission_id"] = report.submission_id

    if _has_model_field(ReportExport, "requested_by"):
        data["requested_by"] = created_by

    if _has_model_field(ReportExport, "created_by"):
        data["created_by"] = created_by

    if _has_model_field(ReportExport, "export_format"):
        data["export_format"] = export_format.upper()

    if _has_model_field(ReportExport, "format"):
        data["format"] = export_format.lower()

    if _has_model_field(ReportExport, "file_name"):
        data["file_name"] = file_name

    if _has_model_field(ReportExport, "stored_path"):
        data["stored_path"] = stored_path

    if _has_model_field(ReportExport, "storage_path"):
        data["storage_path"] = stored_path

    if _has_model_field(ReportExport, "mime_type"):
        data["mime_type"] = mime_type

    if _has_model_field(ReportExport, "size_bytes"):
        data["size_bytes"] = size_bytes

    if _has_model_field(ReportExport, "status"):
        data["status"] = status_value

    if _has_model_field(ReportExport, "error_message"):
        data["error_message"] = error_message

    if _has_model_field(ReportExport, "finished_at"):
        data["finished_at"] = datetime.now(timezone.utc)

    report_export = ReportExport(**data)

    db.add(report_export)
    db.flush()

    return report_export


def _safe_record_audit_log(
    db: Session,
    user_id: UUID | None,
    report: Report,
    export_record: ReportExport,
    export_format: str,
) -> None:
    try:
        from app.services.audit_service import record_audit_log

        record_audit_log(
            db=db,
            user_id=user_id,
            action="EXPORT_REPORT",
            resource_type="report",
            resource_id=str(report.id),
            message=f"Report {export_format.upper()} exported.",
            details={
                "export_id": str(export_record.id),
                "format": export_format.lower(),
            },
        )

    except Exception:
        return


def create_report_export(
    db: Session,
    report: Report,
    export_format: str,
    include_raw_citation: bool = True,
    created_by: UUID | None = None,
) -> ReportExport:
    normalized_format = export_format.strip().lower()

    if normalized_format not in {"pdf", "docx", "xlsx"}:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail={
                "error_code": "EXPORT_FORMAT_NOT_IMPLEMENTED",
                "message": "Hiện tại chỉ hỗ trợ export PDF, DOCX và XLSX.",
                "details": {
                    "format": export_format,
                },
            },
        )

    export_dir = _get_export_root_dir()

    submission = getattr(report, "submission", None)
    owner_label = getattr(submission, "owner_label", None)
    file_owner = _safe_file_part(owner_label, str(report.submission_id))

    file_name = (
        f"trustlens_report_{file_owner}_"
        f"{datetime.now().strftime('%Y%m%d%H%M%S')}.{normalized_format}"
    )

    output_path = export_dir / file_name

    if normalized_format == "pdf":
        mime_type = PDF_MIME_TYPE
    elif normalized_format == "docx":
        mime_type = DOCX_MIME_TYPE
    else:
        mime_type = XLSX_MIME_TYPE

    try:
        if normalized_format == "pdf":
            _write_pdf_file(
                output_path=output_path,
                report=report,
                include_raw_citation=include_raw_citation,
            )
        elif normalized_format == "docx":
            _write_docx_file(
                output_path=output_path,
                report=report,
                include_raw_citation=include_raw_citation,
            )
        else:
            _write_xlsx_file(
                output_path=output_path,
                report=report,
                include_raw_citation=include_raw_citation,
            )

        export_record = _record_export_safely(
            db=db,
            report=report,
            export_format=normalized_format,
            file_name=file_name,
            stored_path=str(output_path),
            mime_type=mime_type,
            status_value="COMPLETED",
            created_by=created_by,
        )

        _safe_record_audit_log(
            db=db,
            user_id=created_by,
            report=report,
            export_record=export_record,
            export_format=normalized_format,
        )

        db.commit()
        db.refresh(export_record)

        return export_record

    except HTTPException:
        db.rollback()
        raise

    except Exception as exc:
        db.rollback()

        failed_record = _record_export_safely(
            db=db,
            report=report,
            export_format=normalized_format,
            file_name=file_name,
            stored_path=str(output_path),
            mime_type=mime_type,
            status_value="FAILED",
            created_by=created_by,
            error_message=str(exc),
        )

        db.commit()
        db.refresh(failed_record)

        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={
                "error_code": f"{normalized_format.upper()}_EXPORT_FAILED",
                "message": f"Không thể export report sang {normalized_format.upper()}.",
                "details": {
                    "report_id": str(report.id),
                    "submission_id": str(report.submission_id),
                    "error": str(exc),
                },
            },
        )


def export_submission_report_to_pdf(
    db: Session,
    submission_id: UUID,
    current_user=None,
) -> ReportExport:
    report = _get_latest_report_by_submission_id(
        db=db,
        submission_id=submission_id,
    )

    return create_report_export(
        db=db,
        report=report,
        export_format="pdf",
        include_raw_citation=True,
        created_by=getattr(current_user, "id", None),
    )


def export_submission_report_to_docx(
    db: Session,
    submission_id: UUID,
    current_user=None,
) -> ReportExport:
    report = _get_latest_report_by_submission_id(
        db=db,
        submission_id=submission_id,
    )

    return create_report_export(
        db=db,
        report=report,
        export_format="docx",
        include_raw_citation=True,
        created_by=getattr(current_user, "id", None),
    )


def export_submission_report_to_xlsx(
    db: Session,
    submission_id: UUID,
    current_user=None,
) -> ReportExport:
    report = _get_latest_report_by_submission_id(
        db=db,
        submission_id=submission_id,
    )

    return create_report_export(
        db=db,
        report=report,
        export_format="xlsx",
        include_raw_citation=True,
        created_by=getattr(current_user, "id", None),
    )

def get_submission_report(
    db: Session,
    submission_id: UUID,
    current_user=None,
) -> Report:
    """
    Lấy report mới nhất của một submission.

    Dùng cho endpoint reports.py.
    Nếu current_user được truyền vào thì kiểm tra quyền truy cập submission.
    """

    if current_user is not None:
        get_accessible_submission_or_404(
            db=db,
            submission_id=submission_id,
            current_user=current_user,
        )

    report = db.execute(
        select(Report)
        .where(Report.submission_id == submission_id)
        .order_by(Report.created_at.desc())
    ).scalars().first()

    if report is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={
                "error_code": "REPORT_NOT_FOUND",
                "message": "Không tìm thấy report của bài nộp này.",
                "details": {
                    "submission_id": str(submission_id),
                },
            },
        )

    return report
